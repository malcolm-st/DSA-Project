#############################################################################################
#######################                                            ##########################
#######################                  IMPORTS                   ##########################
#######################                                            ##########################
#############################################################################################

import tkinter as tk
# import tkinter.ttk as ttk
from tkinter import font, filedialog, messagebox, simpledialog, ttk, Scrollbar
import docx
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import csv
import win32api
import win32gui

# retrieve imports
import git
import os
import subprocess
import shutil
from tqdm import tqdm
from threading import Thread

# update imports
import requests
from bs4 import BeautifulSoup
import zipfile
import io
import time
import threading

# json to csv imports
import json

# for URL openings
import webbrowser

from tkinter import *
from SystemChecker import *
from VendorAnalysis import *
from YearAnalysis import *
from UpdateCVE import *
from RetrieveCVE import *


#############################################################################################
#######################                                            ##########################
#######################       Global Variables and Functions       ##########################
#######################                                            ##########################
#############################################################################################

# Global flag variable (for update)
is_program_running = True

bar_chart_created = False
bar_chart = None
uploaded=False
cache = {}
rows_to_display = []

def update_upload_status():
    # Access the global variable
    global uploaded

    if uploaded:
        uploaded = False
        show_cvesearch_page()
    # Modify the global variable
    else:
        uploaded = True

# Call the function to update the global variable
def get_rows_to_display():
    return rows_to_display

#############################################################################################
#######################                                            ##########################
#######################      "UPDATE" BUTTON RELATED FUNCTION      ##########################
#######################                                            ##########################
#############################################################################################

# Indefinitely check for updates
def check_for_updates():
    # Define the repository details
    username = "CVEProject"
    repository = "cvelistV5"
    branch = "main"

    # Define the API endpoint and headers
    api_endpoint = f"https://api.github.com/repos/{username}/{repository}/branches/{branch}"
    headers = {"Accept": "application/vnd.github.v3+json",
               "Authorization": "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"}

    # Initialize the last commit hash
    last_commit_hash = None

    # Loop indefinitely to check for updates
    while is_program_running:
        # Send a GET request to the API endpoint to get branch details
        response = requests.get(api_endpoint, headers=headers)

        # Check if the response is successful
        if response.ok:
            # Get the latest commit hash
            commit_hash = response.json()["commit"]["sha"]

            print("Checking Repo for commit hash")

            # Check if this is the first loop or if the commit hash has changed
            if last_commit_hash is None or commit_hash != last_commit_hash:
                print("Updating Local Database")
                # Update the local folder
                update_cve()

                # Update the last commit hash
                last_commit_hash = commit_hash

                # Log the update
                print(f"Updated local folder withnew changes from GitHub repository. Commit hash: {commit_hash}")

        else:
            # Log the error
            print(f"Error checking for updates: {response.status_code} {response.reason}")

        # Wait for 1 hour before checking again
        # time.sleep(3600)

        # Updates Database every minute
        time.sleep(60)

# Function to call the check_for_updates function in a separate thread so GUI will not crash
def run_update_checker():
    thread = threading.Thread(target=check_for_updates)
    thread.start()

def update():
    print("Update button clicked")

    # check_for_updates() will update CVEs properly, but GUI will not respond due to while loop
    # Use threading to call the function instead, run_update_checker()
    run_update_checker()

    # Testing updating csv file from .json updates file
    # update_csv_from_json()

#############################################################################################
#######################                                            ##########################
#######################  "Retrieve Data" BUTTON RELATED FUNCTION   ##########################
#######################                                            ##########################
#############################################################################################

# For retrieval
def retrieve_data():
    print("Retrieve Data button clicked")
    # Create a progress bar
    total_steps = 4
    progress_bar = tqdm(total=total_steps, unit="step")

    if not os.path.exists("scrapedCVE"):
        os.makedirs("scrapedCVE")

    # Retrieve CVEs from database
    pull_cves(progress_bar)
    # progress_bar.update(1)

    # Call the function to remove folders and files
    remove_folders_files(local_folder, progress_bar)
    # progress_bar.update(1)

    # Call the function to move files
    move_files("scrapedCVE/cves", "compiledCVE", progress_bar)

    # Removes original scrapedCVE folder
    shutil.rmtree("scrapedCVE")
    # progress_bar.update(1)

    # Convert .json files into a single CVECSV.csv files
    # Provide the folder containing JSON files and the desired CSV file
    dir_path = os.path.dirname(os.path.realpath(__file__))
    json_filename = 'compiledCVE'
    csv_filename = 'CVECSV.csv'
    json_folder = os.path.join(dir_path, json_filename)
    csv_file = os.path.join(dir_path, csv_filename)

    json_to_csv(json_folder, csv_file)

    progress_bar.update(1)

    # Close the progress bar
    progress_bar.close()

    # Signal that the retrieval process is complete
    retrieval_complete.set(True)

def check_retrieval_status():
    if retrieval_complete.get():
        # Show success message
        success_label.config(text="Retrieve Data Successful")
        # Destroy the loading window
        loading_window.destroy()
        
    else:
        # Continue checking the status after 100ms
        root.after(100, check_retrieval_status)

def open_loading_window():
    global loading_window, success_label, retrieval_complete
    loading_window = tk.Toplevel(root)
    loading_window.title("Loading...")
    loading_window.geometry("400x100")
    loading_window.resizable(False, False)

    progress_label = ttk.Label(loading_window, text="Retrieving data...")
    progress_label.pack(pady=10)

    progress_bar = ttk.Progressbar(loading_window, mode="indeterminate", length=350)
    progress_bar.pack(pady=10)
    progress_bar.start()

    success_label = ttk.Label(loading_window, text="")
    success_label.pack(pady=10)

    retrieval_complete = tk.BooleanVar()
    retrieval_complete.set(False)

    # from RetrieveCVE import retrieve_data
    loading_thread = Thread(target=retrieve_data)
    loading_thread.start()

    check_retrieval_status()

#############################################################################################
#######################                                            ##########################
#######################             Vendor Analysis                ##########################
#######################                                            ##########################
#############################################################################################

# Analysis Page
def show_vendor_analysis_page():
    
    print("Analysis Page Clicked")
    
    # Hide label_analysis and analysis widgets
    label_analysis.pack_forget()
    analysis.pack_forget()

    # Ask for number of items
    num_items = simpledialog.askinteger("Number of vendors", "Enter the number of vendors to display (Max 20):")

    # Ensures num_items is between 1 to 20 inclusive. Prompts user to re-key the value again if too large
    while num_items > 20 or num_items < 1:
        num_items = simpledialog.askinteger("Number of vendors", "The value you have entered is not valid.\n\nEnter the number of vendors to display (Max 20):")
    
    if num_items is None:
        return

    for child in page_frame.winfo_children():
        child.pack_forget()
    home_page.pack_forget()
    label_toolcheck.pack_forget()
    label_analysis.pack(pady=20, side="top")
    analysis.pack(fill="both", expand=True)

    if not bar_chart_created:
        create_bar_chart()
    else:
        # Clear bar chart data before replotting
        plt.clf()
        # Remove the bar chart from the frame
        bar_chart.get_tk_widget().pack_forget()
        
    # Repack the bar chart and existing label
    bar_chart.get_tk_widget().pack(fill='both', expand=True)
    label_analysis.pack(pady=20, side="top")

    print("Number items below:")
    print(num_items)
    vendor_frequency_analysis(num_items)

def create_bar_chart():
    global bar_chart_created, bar_chart

    if not bar_chart_created:

        # Display the bar chart
        fig = plt.gcf()  # Get the current figure
        bar_chart = FigureCanvasTkAgg(fig, master=analysis)
        bar_chart.draw()
        bar_chart.get_tk_widget().pack(fill='both', expand=True)
        bar_chart_created = True

#############################################################################################
#######################                                            ##########################
#######################               Year Analysis                ##########################
#######################                                            ##########################
#############################################################################################

def show_year_analysis_page():

    label_year.pack_forget()
    year.pack_forget()

    for child in page_frame.winfo_children():
        child.pack_forget()
    home_page.pack_forget()
    label_toolcheck.pack_forget()
    #label_cvesearch.pack_forget()
    label_year.pack(pady=20, side="top")
    year.pack(fill="both", expand=True)

    create_year_chart(year)
    
#############################################################################################
#######################                                            ##########################
#######################        System Security Checker Page        ##########################
#######################                                            ##########################
#############################################################################################

def show_toolcheck_page():
    
    print("Tool Page clicked")

    for child in page_frame.winfo_children():
        child.pack_forget()

    home_page.pack_forget()
    label_analysis.pack_forget()
    label_toolcheck.pack(pady=20, side="top")
    toolcheck.pack(fill="both", expand=True)

    windows_defender_result = check_windows_defender_settings()
    firewall_result = check_firewall_settings()
    
    label_toolcheck.config(text=f"Windows Defender Status:\n{windows_defender_result}\n\n{firewall_result}", font=("Arial", 40))

#############################################################################################
#######################                                            ##########################
#######################             CVE Search Page                ##########################
#######################                                            ##########################
#############################################################################################

#This allows user to click a displayed CVE output and be brought to the actual CVE webpage
def open_url(event):
    
    selected_row = show_cvesearch_page.results_tree.focus()
    cve_id = show_cvesearch_page.results_tree.item(selected_row)["values"][0]
    url = "https://nvd.nist.gov/vuln/detail/" + cve_id

    confirmed = messagebox.askyesno("Confirmation", "Are you sure you want to open the URL in a web browser?")
    if confirmed:
        webbrowser.open_new(url)

    

#Main function in the CVE Search Page to display all the buttons, search bar, etc.
def show_cvesearch_page():


    # Clear previous contents of the page
    for child in page_frame.winfo_children():
        child.pack_forget()

    home_page.pack_forget()
    label_analysis.pack_forget()
    label_toolcheck.pack_forget()
    cvesearch.pack(fill="both", expand=True)

     # Check if search bar and results frame already exist
    if hasattr(show_cvesearch_page, 'search_frame'):
        show_cvesearch_page.search_frame.pack()
    else:
        label_cvesearch = tk.Label(cvesearch, text="CVE Search", font=title_font, bg="#f2f2f2")
        # Pack and lift the label_cvesearch widget
        label_cvesearch.pack(side="top", pady=20)
        # Create search bar
        show_cvesearch_page.search_frame = tk.Frame(cvesearch, bg="light blue")
        show_cvesearch_page.search_frame.pack(side="top", fill="x")

        # Add search bar
        search_label = tk.Label(show_cvesearch_page.search_frame, text="Search:", font=button_font, bg="light blue")
        search_label.pack(side="left", padx=10)

        show_cvesearch_page.search_entry = tk.Entry(show_cvesearch_page.search_frame, font=button_font, width=30)
        show_cvesearch_page.search_entry.pack(side="left", padx=10)

        # Add sort by dropdown menu
        sort_label = tk.Label(show_cvesearch_page.search_frame, text="Sort by:", font=button_font, bg="light blue")
        sort_label.pack(side="left", padx=10)

        show_cvesearch_page.sort_var = tk.StringVar()
        sort_options = ["CveID Asc", "CveID Desc", "Score Asc", "Score Desc"]
        sort_dropdown = tk.OptionMenu(show_cvesearch_page.search_frame, show_cvesearch_page.sort_var, *sort_options)
        sort_dropdown.pack(side="left", padx=10)

     # remove the export_button from the return_frame
    if hasattr(return_frame, 'export_button'):
        return_frame.export_button.pack_forget()
        del return_frame.export_button

    if hasattr(show_cvesearch_page, 'results_frame'):
        show_cvesearch_page.results_frame.pack()
    else:
        # Create results frame
        show_cvesearch_page.results_frame = tk.Frame(cvesearch)
        show_cvesearch_page.results_frame.pack(fill="both", expand=True)

        # Create a Treeview widget to display the results
        show_cvesearch_page.results_tree = ttk.Treeview(show_cvesearch_page.results_frame, show="headings")
 
        # Create a vertical scrollbar
        show_cvesearch_page.tree_y_scrollbar = ttk.Scrollbar(show_cvesearch_page.results_frame, orient="vertical", command=show_cvesearch_page.results_tree.yview)

        # Create a horizontal scrollbar
        show_cvesearch_page.tree_x_scrollbar = ttk.Scrollbar(show_cvesearch_page.results_frame, orient="horizontal", command=show_cvesearch_page.results_tree.xview)

        # Configure the scrollbar and Treeview
        show_cvesearch_page.results_tree.configure(yscrollcommand=show_cvesearch_page.tree_y_scrollbar.set, xscrollcommand=show_cvesearch_page.tree_x_scrollbar.set)
        show_cvesearch_page.tree_y_scrollbar.pack(side="right", fill="y")

        show_cvesearch_page.results_tree.pack(fill="both", expand=True)
        show_cvesearch_page.tree_x_scrollbar.pack(side="bottom", fill="x")

    # File Path
    csv_filename = 'CVECSV.csv'

    # Read the CSV file and update the data
    with open(csv_filename, "r", encoding="utf-8") as file:
        reader = csv.reader(file)
        data = list(reader)
        
    display_csv_data(data)

    # # Create search button if it doesn't exist
    if not hasattr(show_cvesearch_page, 'search_button'):
        show_cvesearch_page.search_button = tk.Button(show_cvesearch_page.search_frame, text="Search", command=search_cve_wrapper, **button_style)
        show_cvesearch_page.search_button.pack(side=tk.LEFT, padx=10, pady=5)

        show_cvesearch_page.upload_button = tk.Button(show_cvesearch_page.search_frame, text="Upload", command=upload, **button_style)
        show_cvesearch_page.upload_button.pack(side=tk.LEFT,pady=5)

    # Forget and re-pack the search button to ensure it is displayed correctly
    show_cvesearch_page.search_button.pack_forget()
    show_cvesearch_page.search_button.pack(side=tk.LEFT, padx=10, pady=5)
    show_cvesearch_page.upload_button.pack_forget()
    show_cvesearch_page.upload_button.pack(side=tk.LEFT, pady=5)

    
#This function displays the csv data
def display_csv_data(data):

    # Clear existing data in Treeview
    show_cvesearch_page.results_tree.delete(*show_cvesearch_page.results_tree.get_children())

    header = ['CveID', 'Vendor', 'Score', 'Description']
    show_cvesearch_page.results_tree["columns"] = header

    # Configure column names and properties
    column_widths = [10, 10, 10, 400]  # Specify the width for each column
    column_min_widths = [100, 200, 100, 20000]  # Specify the minimum width for each column

    for i, col in enumerate(header):
        show_cvesearch_page.results_tree.heading(col, text=col, anchor=tk.W)
        show_cvesearch_page.results_tree.column(col, width=column_widths[i], minwidth=column_min_widths[i])

    # check if first row contains headers
    if header == data[0]:
        #if first row contains headers then display from second row onwards
        for row in data[1:]:
            show_cvesearch_page.results_tree.insert("", tk.END, values=row)
    else:
        if len(data) == 1:
            for row in data:
                show_cvesearch_page.results_tree.insert("", tk.END, values=row)
        else:
            for row in data:
                show_cvesearch_page.results_tree.insert("", tk.END, values=row)
    
    # This to bind the action of double clicking to the function open_url     
    show_cvesearch_page.results_tree.bind("<Double-1>", open_url)

def search_cve_wrapper():
    search_text = show_cvesearch_page.search_entry.get()
    if uploaded:
        all_search(search_text, get_rows_to_display())
    else:
        all_search(search_text, None)

def all_search(search_query, uploaded_data):
    # Create a cache key using the search query and sort option
    sort_option = show_cvesearch_page.sort_var.get()
    cache_key = (search_query, sort_option)
    global uploaded
    print(uploaded)

    # Check if search query and sort option exist in the cache
    if cache_key in cache and not uploaded:
        print("Retrieving results from cache...")
        results = cache[cache_key]

    else:
        if uploaded_data:
            data = pd.DataFrame(uploaded_data)
        else:
            data = pd.read_csv('CVECSV.csv', encoding='utf-8')

        search_query = str(search_query).lower()
        filtered_data = data[data.apply(lambda row: any(search_query in str(cell).lower() for cell in row), axis=1)]
        results = filtered_data.values.tolist()


        vendor_column_index = 1
        for row in results:
            if pd.isna(row[vendor_column_index]):
                row[vendor_column_index] = "n/a"
            elif row[vendor_column_index] == "":
                row[vendor_column_index] = "n/a"

        score_column_index = 2
        for row in results:
            if pd.isna(row[score_column_index]):
                row[score_column_index] = ""
            elif row[score_column_index] == "":
                row[score_column_index] = ""

        # Sort the results based on the selected option
        if sort_option == "CveID Asc":
            results.sort(key=lambda x: x[0], reverse=False)
        elif sort_option == "CveID Desc":
            results.sort(key=lambda x: x[0], reverse=True)
        elif sort_option == "Score Asc":
            results.sort(
                key=lambda x: float(x[2]) if isinstance(x[2], str) and x[2].replace('.', '', 1).isdigit() else float(
                    'inf') if isinstance(x[2], str) else x[2], reverse=False)
        elif sort_option == "Score Desc":
            results.sort(
                key=lambda x: float(x[2]) if isinstance(x[2], str) and x[2].replace('.', '', 1).isdigit() else float(
                    '-inf') if isinstance(x[2], str) else x[2], reverse=True)



        # Cache the results
        if not uploaded:
            cache[cache_key] = results
            # Update the cache counter
            show_cvesearch_page.cache_counter.config(text="Cache Size: " + str(len(cache)))


        # Check cache size and clear if necessary
        if len(cache) > 100:
            print("Clearing cache...")
            clear_cache()


    display_csv_data(results)

def clear_cache():
    cache.clear()
    show_cvesearch_page.cache_counter.config(text="Cache Size: 0")


# Upload button
def upload():

    # Use the global keyword to access the global rows_to_display variable
    global uploaded
    global rows_to_display  

    file_paths = filedialog.askopenfilenames(filetypes=[("Word Documents", ".docx"), ("Text Files", ".txt")])

    # Reset the rows_to_display list
    rows_to_display = []  

    for file_path in file_paths:

        # Check if the file path ends with .docx or .txt
        if file_path.endswith(".docx") or file_path.endswith(".txt"):
            # Call the search_csv_by_id function and append the results to the list
            rows_to_display += search_csv_by_id(file_path, 'CVECSV.csv')

            # Ran the display_csv_data(rows_to_display) to display all the relevant info related to uploaded data
            display_csv_data(rows_to_display)

            if not uploaded:
                update_upload_status()

            # Create export button if it doesnt exist
            if not hasattr(return_frame, 'export_button'): 
                export_button = tk.Button(return_frame, text="Export", command=lambda: export_to_csv(rows_to_display), **button_style)
                export_button.pack(side=tk.LEFT, padx=10, pady=10, anchor='center')

                # Set the export_button attribute on the return_frame object
                return_frame.export_button = export_button 
    
            else:
                return_frame.export_button.pack()

        # error handling statement
        else:
            print("Invalid file format. Please upload a DOCX or TXT file.")

#############################################################################################
#######################                                            ##########################
#######################       USING DICTIONARY BASED TO SEARCH     ##########################
#######################                                            ##########################
#############################################################################################

def search_csv_by_id(file_path, csv_file):

    # Initialize an empty list to store the IDs
    ids = []
    
    # Check if the file path ends with .txt
    if file_path.endswith(".txt"):

        # Open the file in read mode
        with open(file_path) as f:

            # Read each line, strip it of leading/trailing whitespace, and add it to the list of IDs if it's not empty
            ids = [line.strip() for line in f if line.strip()]

    # Check if the file path ends with .docx
    elif file_path.endswith(".docx"):

        # Open the DOCX file using the docx library
        doc = docx.Document(file_path)

        # Iterate over each paragraph in the document
        for paragraph in doc.paragraphs:
            # Get the text of the paragraph, strip it of leading/trailing whitespace
            line = paragraph.text.strip()
            # If the line is not empty, add it to the list of IDs
            if line:
                ids.append(line)
    else:
        # If the file is not a TXT or DOCX file, print an error message and return an empty list
        print("Invalid file format. Please upload a DOCX or TXT file.")
        return []
    
    # Open the CSV file in read mode with utf-8 encoding
    with open(csv_file, "r", encoding="utf-8") as file:
        # Create a CSV reader object to read the contents of the file
        reader = csv.reader(file)
        
        header = next(reader)
        
        # Identify header with the label 'CveID'
        cve_id_index = header.index('CveID')
        
        data = {row[cve_id_index]: row for row in reader}

    # Initialize an empty list to store the rows to display
    rows_to_display = []
    
    # Initialize an empty set to store unique CVE IDs which will be used to avoid displaying duplicate rows
    unique_cve_ids = set()
    
    # Iterate over each ID in the list of IDs
    for cve_id in ids:

        # Check if the ID is in the data dictionary and has not already been added to the unique_cve_ids set
        if cve_id in data and cve_id not in unique_cve_ids:

            # Add the row corresponding to that ID to the rows_to_display list
            rows_to_display.append(data[cve_id])

            # Add the ID to the unique_cve_ids set
            unique_cve_ids.add(cve_id)

    # Check if no rows were found or if the last row added does not match the last ID searched for
    if not rows_to_display or rows_to_display[-1][cve_id_index] != cve_id:
        
        # Print a message indicating that no results were found for that ID
        print(f"No results found for {cve_id}")

    # Return the list of rows to display
    return rows_to_display

#############################################################################################
#######################                                            ##########################
#######################         USING BINARY SEARCH TO SEARCH      ##########################
#######################                                            ##########################
#############################################################################################
#To use the binary search function instead just highlight the whole thing and uncomment. However, remember to comment the dictionary key search function above.
#This is solely for project assessment purpose, it is not used but is left for grading purposes to show the difference in both codes.
#Look at report for more details

# def search_csv_by_id(file_path, csv_file):
    
#     global rows_to_display
    
#     # reset the rows_to_display list    
#     rows_to_display = []  

#     # Initialize an empty list to store the IDs
#     ids = []
    
#     # Check if the file path ends with .txt
#     if file_path.endswith(".txt"):

#         # Open the file in read mode
#         with open(file_path) as f:

#             # Read each line, strip it of leading/trailing whitespace, and add it to the list of IDs if it's not empty
#             ids = [line.strip() for line in f if line.strip()]

#     # Check if the file path ends with .docx
#     elif file_path.endswith(".docx"):

#         # Open the DOCX file using the docx library
#         doc = docx.Document(file_path)

#         # Iterate over each paragraph in the document
#         for paragraph in doc.paragraphs:

#             # Get the text of the paragraph, strip it of leading/trailing whitespace
#             line = paragraph.text.strip()

#             # If the line is not empty, add it to the list of IDs
#             if line:
#                 ids.append(line)
#     else:
#         # If the file is not a TXT or DOCX file, print an error message and return an empty list
#         print("Invalid file format. Please upload a DOCX or TXT file.")
#         return []
    
#     # Open the CSV file in read mode with utf-8 encoding
#     with open(csv_file, "r", encoding="utf-8") as file:
        
#         # Create a CSV reader object to read the contents of the file
#         reader = csv.reader(file)
        
#         header = next(reader)
        
#         # Identify header with the label 'CveID'
#         cve_id_index = header.index('CveID')
        
#         data = sorted([row for row in reader], key=lambda x: x[cve_id_index])
    
#     #Binary search function
#     def binary_search(data, target):
#         # Set the initial values for the low and high indices
#         low = 0
#         high = len(data) - 1
        
#         # Continue searching while the low index is less than or equal to the high index
#         while low <= high:
#             # Calculate the middle index
#             mid = (low + high) // 2
#             # Check if the value at the middle index is equal to the target value
#             if data[mid][cve_id_index] == target:
#                 # If it is, return the row at the middle index
#                 return data[mid]
#             # If the value at the middle index is less than the target value
#             elif data[mid][cve_id_index] < target:
#                 # Set the new low index to be one more than the middle index
#                 low = mid + 1
#             # If the value at the middle index is greater than the target value
#             else:
#                 # Set the new high index to be one less than the middle index
#                 high = mid - 1
        
#         # If the target value is not found, return None
#         return None
    
#     # Initialize an empty set to store unique CVE IDs (to avoid displaying duplicate rows)
#     unique_cve_ids = set()
    
#     # Iterate over each ID in the list of IDs
#     for cve_id in ids:
#         row = binary_search(data, cve_id)
        
#         if row and cve_id not in unique_cve_ids:
#             rows_to_display.append(row)
#             unique_cve_ids.add(cve_id)

#     # Remove any duplicated output
#     rows_to_display = list(set(map(tuple, rows_to_display)))  

#     # Return the list of rows to display
#     return rows_to_display

#function used for exporting the displayed rows
def export_to_csv(rows_to_display):

    if not rows_to_display:
        print("No data to save. Please ensure that the fields are outputted before exporting.")
        return

    # Open a save file dialog
    file_path = filedialog.asksaveasfilename(defaultextension=".csv", filetypes=[("CSV Files", "*.csv")])
    
    if not file_path:
        # The user cancelled the dialog
        return
    
    # Open the selected file in write mode
    with open(file_path, 'w', newline='') as csvfile:
        # Create a CSV writer object
        writer = csv.writer(csvfile)
        
        # Write the header row
        writer.writerow(["CveID", "Vendor", "Score", "Description"])  
        
        # Write the rows
        for row in rows_to_display:
            writer.writerow(row)
    
    print(f"Data has been successfully saved to {file_path}.")

#############################################################################################
#######################                                            ##########################
#######################                 Home Page                  ##########################
#######################                                            ##########################
#############################################################################################

def show_home_page():
    
    toolcheck.pack_forget()
    cvesearch.pack_forget()
    analysis.pack_forget()
    year.pack_forget()
    home_page.pack(fill="both", expand=True)   

#############################################################################################
#######################                                            ##########################
#######################     GUI Button/Label/Text Creation Codes   ##########################
#######################                                            ##########################
#############################################################################################

root = tk.Tk()
root.title("CVE Aggregator")

# Get the screen dimensions
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

# # Calculate the window position to appear right above the taskbar
# taskbar_hwnd = win32gui.FindWindow("Shell_TrayWnd", None)
# taskbar_rect = win32gui.GetWindowRect(taskbar_hwnd)
# taskbar_height = taskbar_rect[1] - taskbar_rect[3]
# window_x = 0
# window_y = screen_height - (screen_height-taskbar_height)

# Set the window size and position using geometry()
# root.geometry(f"{screen_width}x{screen_height-taskbar_height}+{window_x}+{window_y}")

# Calculate the font size based on the screen height
title_font_size = int(screen_height / 20)
button_font_size = int(screen_height / 40)

# Create custom fonts with the calculated sizes
title_font = font.Font(size=title_font_size, weight="bold")
button_font = font.Font(size=button_font_size)

# Create a Frame to hold the pages
page_frame = tk.Frame(root)
page_frame.pack(fill="both", expand=True)

# Create the Home page
home_page = tk.Frame(page_frame, bg="#f2f2f2")
home_page.pack(fill="both", expand=True)

welcome_label = tk.Label(home_page, text="Welcome to CVE Aggregator", font=title_font, wraplength=screen_width - 100,  bg="#f2f2f2")
welcome_label.pack(pady=20)

# Configure button styles
button_style = {
    "bg": "#4CAF50",
    "fg": "white",
    "activebackground": "#45a049",
    "activeforeground": "white",
    "bd": 0,
    "width": 20,
    "font": button_font,
    "pady": 10
}

update_button = tk.Button(home_page, text="Update", command=update, **button_style)
update_button.pack(pady=10)

retrieve_button = tk.Button(home_page, text="Retrieve Data", command=open_loading_window, **button_style)
retrieve_button.pack(pady=10)

# Create the Page 1
toolcheck = tk.Frame(page_frame, bg="light blue")
label_toolcheck = tk.Label(toolcheck, text="System Security Checker", font=title_font, bg="#f2f2f2")

# Create the Page 2
cvesearch = tk.Frame(page_frame, bg="light blue")

# Create the Page 3
analysis = tk.Frame(page_frame, bg="light blue")
label_analysis = tk.Label(analysis, text="Vendor Analysis", font=title_font, bg="#f2f2f2")

# Create year frame
year = tk.Frame(page_frame, bg="light blue")
label_year = tk.Label(year, text="Year Analysis", font=title_font, bg="#f2f2f2")

# Create navigation buttons
nav_frame = tk.Frame(root, bg="#f2f2f2")
nav_frame.pack(side="bottom", pady=10)

home_button = tk.Button(nav_frame, text="Home", command=show_home_page, **button_style)
home_button.pack(side="left", padx=10)

page1_button = tk.Button(nav_frame, text="System Security Checker", command=show_toolcheck_page, **button_style)
page1_button.pack(side="left", padx=10)

page2_button = tk.Button(nav_frame, text="CVE Search", command=show_cvesearch_page, **button_style)
page2_button.pack(side="left", padx=10)

page3_button = tk.Button(home_page, text="Vendor Analysis", command=show_vendor_analysis_page, **button_style)
page3_button.pack(pady=10)

page4_button = tk.Button(home_page, text="Year Analysis", command=show_year_analysis_page, **button_style)
page4_button.pack(pady=10)

root.geometry(f"{screen_width}x{screen_height}")  # Set window size to full screen

cvesearch = tk.Frame(page_frame, bg="light blue")

# Create a Frame for the return button
return_frame = tk.Frame(cvesearch, bg="light blue")
return_frame.pack(side="bottom", fill="x")
    
return_button = tk.Button(return_frame, text="Reset Filters", command=update_upload_status, **button_style)
return_button.pack(side=tk.LEFT, padx=10, pady=10, anchor='center')

return_button = tk.Button(return_frame, text="Clear Search Cache", command=clear_cache, **button_style)
return_button.pack(side=tk.LEFT, padx=10, pady=10, anchor='center')

# Create a label widget for the cache counter
show_cvesearch_page.cache_counter = tk.Label(return_frame, text="Cache Size: 0", font=button_font, bg="#f2f2f2")
show_cvesearch_page.cache_counter.pack(side="right", padx=15, pady=10)

#############################################################################################
#######################                                            ##########################
#######################           TKINTER MAIN FUNCTIONS           ##########################
#######################                                            ##########################
#############################################################################################

def show_page(page):
    page.tkraise()

# Set initial page
show_page(home_page)

def quit_program():
    root.quit()
    global is_program_running
    is_program_running = False
    root.destroy()

# Bind the on_close function to the close event of the GUI window
root.protocol("WM_DELETE_WINDOW", quit_program)

quit_button = tk.Button(nav_frame, text="Quit", command=quit_program, **button_style)
quit_button.pack(side="left", padx=10)

root.mainloop()
